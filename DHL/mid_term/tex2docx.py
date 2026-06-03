"""
tex2docx.py — 将 mid_term_report.tex 转为格式保持的 DOCX

三步流程：
1. 用 python-docx 创建匹配 ctexart 样式的 reference.docx 模板
2. 调用 pandoc + reference.docx 执行转换
3. 用 python-docx 后处理微调（表格线型、字体确认等）

用法：python tex2docx.py
"""

import os
import subprocess
import sys
import io

# Windows GBK 控制台兼容
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 路径配置 ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEX_FILE = os.path.join(BASE_DIR, "mid_term_report.tex")
DOCX_OUT = os.path.join(BASE_DIR, "mid_term_report.docx")
REF_DOCX = os.path.join(BASE_DIR, "reference.docx")
FIGURES_DIR = os.path.join(BASE_DIR, "figures")


# ════════════════════════════════════════════════════════════════════
# Step 1: 创建 reference.docx 模板
# ════════════════════════════════════════════════════════════════════

def create_reference_docx():
    """生成匹配 ctexart (A4, 12pt, 宋体/黑体) 的 Word 模板"""

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 样式设置 ──
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"          # 西文
    style_normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 中文
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.3
    style_normal.paragraph_format.space_after = Pt(0)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.first_line_indent = Cm(0.74)  # 约两字符缩进

    # Heading 1 → \section（一、二、三、四）
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(16)  # 三号
    h1.font.bold = True
    h1.font.color.rgb = None  # 继承正文颜色（黑色）
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.3
    h1.paragraph_format.first_line_indent = None  # 标题不缩进
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 → \subsection（1.1, 1.2, ...）
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(14)  # 四号
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.3
    h2.paragraph_format.first_line_indent = None
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Caption → 图注/表注
    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(10.5)  # 五号
    caption.font.bold = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.first_line_indent = None

    # Quote / Block Text → 用于 itemize
    quote = doc.styles["Quote"]
    quote.font.name = "Times New Roman"
    quote.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    quote.font.size = Pt(12)
    quote.paragraph_format.left_indent = Cm(1.0)

    # List Paragraph → itemize 列表项
    lp = doc.styles["List Paragraph"]
    lp.font.name = "Times New Roman"
    lp.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    lp.font.size = Pt(12)

    doc.save(REF_DOCX)
    print(f"[Step 1] reference.docx 已生成: {REF_DOCX}")


# ════════════════════════════════════════════════════════════════════
# Step 2: pandoc 转换
# ════════════════════════════════════════════════════════════════════

def run_pandoc():
    """用 pandoc + reference.docx 转换"""

    cmd = [
        "pandoc",
        TEX_FILE,
        "-o", DOCX_OUT,
        "--from", "latex",
        "--to", "docx",
        f"--reference-doc={REF_DOCX}",
        f"--resource-path={FIGURES_DIR}",
    ]

    print(f"[Step 2] pandoc 转换...")
    print(f"  命令: {' '.join(cmd)}")

    result = subprocess.run(cmd, capture_output=True, text=True, cwd=BASE_DIR)

    if result.returncode != 0:
        print(f"  ❌ pandoc 错误:\n{result.stderr}")
        sys.exit(1)

    if result.stderr:
        print(f"  ⚠ pandoc 警告:\n{result.stderr}")

    size_mb = os.path.getsize(DOCX_OUT) / (1024 * 1024)
    print(f"  ✅ 输出: {DOCX_OUT} ({size_mb:.1f} MB)")


# ════════════════════════════════════════════════════════════════════
# Step 3: python-docx 后处理
# ════════════════════════════════════════════════════════════════════

def _is_chinese(text):
    """检测字符串是否含中文"""
    return any('一' <= c <= '鿿' for c in text)


def _set_run_fonts(run, east_asia="宋体", ascii_font="Times New Roman"):
    """强制设置 run 的 eastAsia 和 ascii 字体"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")}'
            f' w:eastAsia="{east_asia}"'
            f' w:ascii="{ascii_font}"'
            f' w:hAnsi="{ascii_font}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), east_asia)
        rFonts.set(qn("w:ascii"), ascii_font)
        rFonts.set(qn("w:hAnsi"), ascii_font)


def _is_caption_text(text):
    """检测段落是否为图注/表注（pandoc 去掉了编号前缀，需按内容特征识别）"""
    import re
    t = text.strip()

    # 排除：正文引用（"图 N 展示了..."、"图 N 所示..."）
    if re.match(r'^图\s*\d+\s*(展示了|所示|给出)', t):
        return False
    # 排除：小节编号标题（"4.5 时间安排"）
    if re.match(r'^\d+\.\d+\s', t):
        return False

    # 显式 "图 N" / "表 N" 前缀
    if re.match(r'^(图|表|Figure|Table)\s*\d', t):
        return True
    # 图注特征：(a)/(b)/(c) 子面板描述
    if re.search(r'（[a-z]）', t) and ('引自' in t or len(t) > 50):
        return True
    # 典型表注关键词（短段落）
    if any(kw in t for kw in ('技术指标', '时间安排')) and len(t) < 30:
        return True
    # 含 "引自" 的短段落
    if '引自' in t and len(t) < 150:
        return True
    return False


def postprocess():
    """后处理：字体、图注样式、表格线型、图片尺寸"""

    doc = Document(DOCX_OUT)
    font_fixes = 0
    caption_fixes = 0
    table_fixes = 0
    img_count = 0

    # ── 3a. 全局字体修正 + 图注识别 ──
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text

        # 识别图注/表注段落 → 改为 Caption 样式
        if _is_caption_text(text) and style_name != "Caption":
            try:
                para.style = doc.styles["Caption"]
            except KeyError:
                pass
            # 居中、五号字、无首行缩进
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = None
            for run in para.runs:
                run.font.size = Pt(10.5)
            caption_fixes += 1

        # 标题不缩进
        if style_name in ("Heading 1", "Heading 2"):
            para.paragraph_format.first_line_indent = None

        # 所有含中文的 run 强制设 eastAsia 字体
        for run in para.runs:
            if _is_chinese(run.text):
                _set_run_fonts(run, east_asia="宋体", ascii_font="Times New Roman")
                font_fixes += 1
            elif run.text.strip():  # 纯英文 run 也要设字体
                _set_run_fonts(run, east_asia="宋体", ascii_font="Times New Roman")

    # ── 3b. 表格线型：模拟 booktabs ──
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 清除所有默认边框
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is not None:
            tblPr.remove(borders)

        # 添加 booktabs 风格：顶线粗、表头下线细、底线粗
        borders_xml = f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tblBorders>'''
        tblPr.append(parse_xml(borders_xml))

        # 表头行下方加细线
        if len(table.rows) >= 2:
            header_row = table.rows[0]
            for cell in header_row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                </w:tcBorders>''')
                tcPr.append(tcBorders)

        # 设置表格内字体
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.first_line_indent = None
                    for run in para.runs:
                        run.font.size = Pt(10.5)
                        _set_run_fonts(run, east_asia="宋体", ascii_font="Times New Roman")

        table_fixes += 1

    # ── 3c. 图片尺寸：确保宽度不超过页面的 85% ──
    # A4 可用宽度 = 21 - 2.5*2 = 16cm, 85% ≈ 13.6cm
    max_width = Cm(13.6)
    for para in doc.paragraphs:
        for run in para.runs:
            el = run._element
            drawings = el.findall(qn("w:drawing"))
            for drawing in drawings:
                for tag in (qn("wp:inline"), qn("wp:anchor")):
                    for cand in drawing.findall(tag):
                        extent = cand.find(qn("wp:extent"))
                        if extent is not None:
                            cx = int(extent.get("cx", "0"))
                            if cx > 0:
                                img_count += 1
                            if cx > max_width:
                                cy = int(extent.get("cy", "0"))
                                ratio = cy / cx if cx > 0 else 1
                                extent.set("cx", str(int(max_width)))
                                extent.set("cy", str(int(max_width * ratio)))

    doc.save(DOCX_OUT)

    size_mb = os.path.getsize(DOCX_OUT) / (1024 * 1024)
    print(f"[Step 3] 后处理完成:")
    print(f"  字体修正: {font_fixes} 个 run")
    print(f"  图注识别: {caption_fixes} 个")
    print(f"  表格线型: {table_fixes} 个")
    print(f"  图片处理: {img_count} 张")
    print(f"  输出: {DOCX_OUT} ({size_mb:.1f} MB)")


# ════════════════════════════════════════════════════════════════════
# 主流程
# ════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("tex2docx: LaTeX → DOCX 格式保持转换")
    print("=" * 60)

    # 检查依赖
    if not os.path.exists(TEX_FILE):
        print(f"❌ 源文件不存在: {TEX_FILE}")
        sys.exit(1)

    # Step 1
    create_reference_docx()

    # Step 2
    run_pandoc()

    # Step 3
    postprocess()

    print(f"\n{'=' * 60}")
    print(f"✅ 转换完成: {DOCX_OUT}")
    print(f"{'=' * 60}")
