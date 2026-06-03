# -*- coding: utf-8 -*-
"""
Generate Supplemental Material .docx for:
"小孔极端约束光场的量子化"
Structure follows the derivation flowchart.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    hs = doc.styles['Heading %d' % level]
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True

# ── Helper functions ──
def add_eq(eq_text, eq_num=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run(eq_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    if eq_num:
        run2 = p.add_run('    (%s)' % eq_num)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ═══════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════
title = doc.add_heading('Supplemental Material', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_st = subtitle.add_run(
    'Derivation Details for "Quantization of Optical Fields '
    'under Extreme Subwavelength Confinement by Small Holes"'
)
run_st.font.name = 'Times New Roman'
run_st.font.size = Pt(12)
run_st.italic = True

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# I. OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_heading('I. 推导框架总览', level=1)

add_body(
    '正文的推导链条可以用一张流程图完整概括。流程图从上方的理论基础出发，'
    '分为三条并行的推导路径：左侧是库仑规范量子化，中间是洛伦兹规范量子化，'
    '右侧是最终物理结果。三条路径在底部通过动力学环节汇合，'
    '最终产出相干态振幅和经典极限下的 Bethe 辐射截面。'
)

add_body(
    '本补充说明的任务是沿着流程图的每一条连线，逐步交代清楚三件事：'
    '（1）每一步在求解什么物理量；'
    '（2）为什么需要求解这个量——它在整个量子化链路中承担什么角色；'
    '（3）该步的结果如何自然地引出下一步。'
    '读完这份材料之后，读者应当能从头到尾复述出从经典偶极模型到量子化场算符、'
    '再到经典极限回归的完整逻辑链。'
)

# ═══════════════════════════════════════════════════════════════
# II. COULOMB GAUGE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('II. 库仑规范下的正则量子化', level=1)
add_body('【对应流程图：左分支，从 Lagrangian (L) 到 A+ 的 Weyl 角谱展开】')

# II.A
doc.add_heading('II.A 为什么从拉格朗日量出发？', level=2)

add_body(
    '经典场论的量子化没有唯一入口，但正则量子化的标准路线要求先写出系统的拉格朗日量密度 L。'
    '原因在于，L 直接决定广义坐标的选取和正则共轭动量的定义——'
    '而这两个对象正是量子化时需要提升为算符的东西。'
    '跳过拉格朗日量、直接写对易关系，会丧失对约束条件（如库仑规范条件 div A = 0）的系统处理能力。'
)

add_body(
    '库仑规范下自由电磁场的拉格朗日量密度为正文 Eq. 4：'
)

add_eq('L = (eps_0/2) A_dot^2 - (1/2 mu_0)(curl A)^2', 'S1')

add_body(
    '第一项是场的"动能"（时间导数的平方），第二项是"势能"（空间导数的平方）。'
    '对 L 施加 Euler-Lagrange 方程，可以复现经典的亥姆霍兹波动方程。'
    '这一验证至关重要：如果 L 的选取无法回到经典运动方程，后续量子化就建立在错误的基础上。'
)

# II.B
doc.add_heading('II.B 矢量势与共轭动量：从经典到量子的桥梁', level=2)

add_body(
    '在拉格朗日力学框架中，广义坐标选为矢量势 A(r, t)，'
    '其正则共轭动量定义为 L 对 A_dot 的偏导。正文 Eq. 5 直接给出：'
)

add_eq('Pi(r, t) = dL/dA_dot = eps_0 A_dot = -eps_0 E', 'S2')

add_body(
    '这个结果揭示了一个重要的物理对应：在库仑规范中，'
    '磁矢势的共轭动量恰好等于负的电场乘以真空介电常数。'
    '量子化的核心操作——将经典变量提升为算符——因此等价于同时将 A 和 E 提升为算符，'
    '并规定它们之间的对易关系。这就是下一步要解决的问题。'
)

# II.C
doc.add_heading('II.C 三维模展开：为什么需要完备基？', level=2)

add_body(
    '量子化之后，场算符 A_hat 是一个算符值函数，无法直接做数值计算。'
    '需要将其展开为一组完备基的线性叠加，'
    '使所有物理信息编码到展开系数（即产生和湮灭算符）中。'
    '正文选择了平面波基，因为自由空间中平面波是亥姆霍兹方程的本征解。'
    '正文 Eq. 6 给出的展开式为：'
)

add_eq(
    'A_hat(r, t) = int d^3k sum_{lam=1,2} N_k [a_hat_{k,lam} eps_{k,lam} e^{-i(w_k t - k.r)} + h.c.]',
    'S3'
)

add_body(
    '这里有两个关键约束来自库仑规范。第一，极化指标 lam 只取 1 和 2（两个横向自由度），'
    '因为 div A = 0 排除了纵向分量。第二，极化矢量 eps_{k,lam} 与波矢 k 正交，'
    '确保每个平面波分量都是纯粹的横波。'
    '产生和湮灭算符 a_dag、a_hat 满足标准的玻色对易关系（正文 Eq. 7），'
    '这是光子作为玻色子的数学表述。'
)

add_body(
    'N_k 是待定归一化常数。它的值不能随意选取——必须由下一步的等时对易关系唯一确定。'
)

# II.D
doc.add_heading('II.D 等时对易关系与归一化系数的确定', level=2)

add_body(
    '正则量子化的数学实质，是规定广义坐标算符与其共轭动量算符在等时刻下满足一个特定的对易关系。'
    '对于电磁场，这个关系是正文 Eq. 8：'
)

add_eq('[A_hat_i(r, t), Pi_hat_j(r\', t)] = i hbar delta_ij_perp(r - r\')', 'S4')

add_body(
    '等号右边的 delta_ij_perp 是横向 delta 函数，而非普通的 delta_ij delta^3(r - r\')。'
    '这一区别的根源在于库仑规范条件 div A = 0：'
    '它把场的自由度限制在横向子空间，因此对易关系也必须只投影到横向。'
    '如果错误地使用普通 delta 函数，计算出的场能量将包含一个无穷大的纵向贡献，与实验矛盾。'
)

add_body(
    '将模展开式 (S3) 和玻色对易关系代入 (S4)，只有 a a_dag 与 a_dag a 的交叉项存活。'
    '利用极化完备性求和（正文 Eq. 9）：'
)

add_eq(
    'sum_lam eps_i(k,lam) eps_j(k,lam) = delta_ij - k_i_hat k_j_hat / |k|^2 = delta_ij_perp(k)',
    'S5'
)

add_body(
    '以及 Fourier 变换恒等式（正文 Eq. 10），可以提取出归一化系数 N_k 必须满足的条件。'
    '最终解出（正文 Eq. 12）：'
)

add_eq('N_k = sqrt[ hbar / (2 eps_0 omega_k (2pi)^3) ]', 'S6')

add_body(
    '这就是连续体归一化下的标准结果。它的物理含义是：'
    '每个模式 (k, lam) 携带的零点能量密度恰好等于 hbar omega_k / 2。'
    '归一化系数一旦确定，场算符的所有性质就完全确定了。'
)

# II.E
doc.add_heading('II.E 场算符与自由哈密顿量', level=2)

add_body(
    '有了 N_k 之后，可以将电场和磁场算符显式写出（正文 Eqs. 13-14）。'
    '电场算符 E_hat 正比于 a_hat（湮灭）和 a_dag（产生）的线性组合，'
    '磁场算符 B_hat 同理但多了一个 k x eps 的叉乘因子——'
    '后者正是磁场"旋度"属性的体现。'
)

add_body(
    '将场算符代入经典哈密顿量的表达式 H = int (eps_0 E^2/2 + B^2/2mu_0) d^3r，'
    '经过对易代数化简，交叉项精确抵消，剩下的只有对角项（正文 Eq. 15）：'
)

add_eq(
    'H_hat_0 = int d^3k sum_{lam=1,2} hbar omega_k (a_dag_{k,lam} a_hat_{k,lam} + 1/2)',
    'S7'
)

add_body(
    '这是相互独立的简谐振子的能量之和。每个模式 (k, lam) 对应一个频率为 omega_k 的量子谐振子，'
    '其能量本征值为 (n_{k,lam} + 1/2) hbar omega_k。零点能不影响动力学过程，后续讨论中略去。'
)

# II.F
doc.add_heading('II.F Weyl 角谱展开：从三维到二维的降维', level=2)

add_body(
    '在实际的小孔衍射问题中，源是单频的（角频率 omega_0），'
    '而且几何结构在 z = 0 平面上具有天然的反射对称性。'
    '这两个条件使得我们可以把三维波矢积分 d^3k 压缩为二维面内波矢积分 d^2k_perp。'
    '这一降维操作在数学上对应 Weyl 角谱展开。'
    '其物理动机是：z 方向的波矢分量 k_z 不再是独立变量，'
    '而是由色散关系 k_z = sqrt(k_0^2 - k_perp^2) 唯一确定——'
    '对于传播模，k_z 为实数；对于倏逝模（k_perp > k_0），k_z 为纯虚数 i kappa。'
)

add_body(
    '正文 Eq. 31 给出的 Sommerfeld 恒等式是这一降维的数学基础。利用它，'
    '矢量势的正频部分可以写成正文 Eq. 32 的形式：'
)

add_eq(
    'A_hat^+(r, t) = int d^2k_perp sum_lam C_{k_perp} a_hat(k_perp, lam) eps_{k_perp,lam} e^{-i(omega_0 t - k.r)}',
    'S8'
)

add_body(
    '其中 C_{k_perp} 是新的归一化振幅。注意到这里 omega_0 替代了 omega_k——'
    '这正是单频条件带来的简化。C_{k_perp} 的确定需要回到真空相关函数，这部分在正文的 Sec. V 中处理。'
)

# ═══════════════════════════════════════════════════════════════
# III. LORENTZ GAUGE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('III. 洛伦兹规范下的量子化', level=1)
add_body('【对应流程图：中间分支，从 Fermi Lagrangian 到规范不变性验证】')

# III.A
doc.add_heading('III.A 为什么需要第二条量子化路径？', level=2)

add_body(
    '库仑规范在非相对论量子光学计算中简洁高效，但它破坏了明显的洛伦兹协变性。'
    '这意味着，如果需要处理相对论修正、圈图计算，或者需要在形式上证明规范不变性，'
    '库仑规范的框架会变得笨拙甚至不可用。洛伦兹规范通过将规范条件提升为协变标量方程（正文 Eq. 16），'
    '保留了理论在洛伦兹变换下的对称性。'
)

add_body(
    '两条路径之间的关系不是二选一，而是互相验证：'
    '如果两种完全不同的量子化方案给出的物理可观测量一致，'
    '那这个结果就是规范无关的——一个真正的物理预测，而非规范选取的数学产物。'
)

# III.B
doc.add_heading('III.B 费米拉格朗日量与四分量体系', level=2)

add_body(
    '洛伦兹规范下的拉格朗日量密度采用费米形式（正文 Eq. 18）：'
)

add_eq('L_F = -(1/4 mu_0) F_{mu nu} F^{mu nu} - (1/2 mu_0) (d_mu A^mu)^2', 'S9')

add_body(
    '第一项是标准 Maxwell 拉格朗日量，第二项是规范固定项——'
    '它的作用是在运动方程层面强制洛伦兹条件 d_mu A^mu = 0。'
    '对 L_F 应用 Euler-Lagrange 方程直接给出协变达朗贝尔方程 dAlembert A^mu = mu_0 J^mu（正文 Eq. 17），'
    '四个分量 A^mu = (phi/c, A) 满足形式上完全相同的波动方程。'
)

add_body(
    '与库仑规范的关键区别在于共轭动量：库仑规范中时间分量动量 Pi^0 = 0（phi 被完全消去），'
    '而洛伦兹规范中 Pi^0 不为零。四个分量地位平等，'
    '代价是引入了两个非物理极化自由度（标量光子 lam=0 和纵向光子 lam=3），'
    '它们对应的态具有负范数——这在物理上不可接受，必须通过下一步的 Gupta-Bleuler 条件来解决。'
)

# III.C
doc.add_heading('III.C Gupta-Bleuler 条件：剔除非物理自由度', level=2)

add_body(
    'Gupta 和 Bleuler 的解决方案不是直接禁止标量光子和纵向光子的存在，'
    '而是规定物理态必须满足一个弱条件（正文 Eq. 24）：'
)

add_eq('(d_mu A_hat^mu)^(+) |phys> = 0', 'S10')

add_body(
    '这里上标 (+) 表示只取正频部分（含湮灭算符）。在动量空间中，'
    '这个条件等价于标量湮灭算符与纵向湮灭算符的特定组合湮灭物理态（正文 Eq. 25）：'
)

add_eq('[a_hat^(0)(k) - a_hat^(3)(k)] |phys> = 0', 'S11')

add_body(
    '这个约束有三个关键后果。第一，物理态的范数总是非负的（正半定性），排除了负范数态的干扰。'
    '第二，标量光子和纵向光子必须成对出现，数目相等——它们的贡献在对易子中精确抵消。'
    '第三，纯横向 Fock 态（即库仑规范所描述的态）自动满足这个条件。'
    '这提示我们，两种规范在物理空间中应该给出相同结果。'
)

# III.D
doc.add_heading('III.D 物理可观测量：为什么两种规范等价？', level=2)

add_body(
    '电场和磁场由规范不变量 F_{mu nu} 决定，因此不依赖于具体规范选取。'
    '但在洛伦兹规范中，场算符包含四个极化分量的贡献，而非仅两个横向分量。'
    '关键验证是：在物理态上，标量光子（lam=0）对 B 的贡献为零（因为 k x eps^(0) = 0），'
    '纵向光子（lam=3）对 B 的贡献也为零（因为 k x k_hat = 0）。'
)

add_body(
    '对于电场，Gupta-Bleuler 条件保证标量与纵向贡献在物理矩阵元之间精确抵消。'
    '因此在物理态之间的电场算符期望值与库仑规范完全一致（正文 Eq. 26）：'
)

add_eq(
    'E_hat_phys^(+) = i int d^3k sum_{lam=1,2} omega_k N_k eps^(lam) a_hat^(lam)(k) e^{i(k.r - omega_k t)}',
    'S12'
)

add_body(
    '同理，哈密顿量在物理态上简化为只含横向模式的形式（正文 Eq. 27），'
    '与库仑规范的 H_hat_0（S7）完全一致。'
    '至此，流程图中间分支的终点与左分支汇合：'
    '两条独立的量子化路径在物理 Hilbert 空间中给出完全相同的结果。'
)

# III.E
doc.add_heading('III.E 倏逝模的 Gupta-Bleuler 条件', level=2)

add_body(
    '小孔近场中大量模式满足 k_perp > k_0，其 z 方向波矢为纯虚数 k_z = i kappa。'
    '对于这些倏逝模式，Gupta-Bleuler 条件的形式略有修正（正文 Eq. 29）：'
)

add_eq(
    '[(omega/c) a_hat^(0)(k) - sqrt(k_perp^2 - kappa^2) a_hat^(3)(k)] |phys> = 0',
    'S13'
)

add_body(
    '这个等式表明，即使对于 z 方向波矢为纯虚数的倏逝模，'
    '标量与纵向光子的相消关系依然成立。'
    '这意味着洛伦兹规范和库仑规范在包含倏逝模的近场区域仍然严格等价。'
    '这也是后续推导中可以放心使用库仑规范的前提——'
    '它不仅适用于传播模，也适用于倏逝模。'
)

# ═══════════════════════════════════════════════════════════════
# IV. DYNAMICS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('IV. 模式归一化与含时动力学', level=1)
add_body('【对应流程图：底部分支，从 Mode Normalization (C_k) 到 Heisenberg Equation】')

# IV.A
doc.add_heading('IV.A 角谱归一化：真空相关函数的匹配', level=2)

add_body(
    'Weyl 角谱展开引入了新的归一化振幅 C_{k_perp}，'
    '它不能从三维的 N_k 直接继承，因为维度已经从三维降到了二维。'
    '确定 C_{k_perp} 的方法是：计算角谱展开下场算符的真空相关函数（正文 Eq. 33-34），'
    '然后与已知的经典 Helmholtz 格林函数在小孔平面（z=0）处的值进行比较。'
)

add_body(
    '匹配的结果给出 |C_{k_perp}|^2 正比于 1/k_z（正文 Eq. 35）。'
    '结合三维对易关系固定的整体归一化，最终得到（正文 Eq. 36）：'
)

add_eq('C_{k_perp} = sqrt[ hbar / (2 eps_0 omega_0 (2pi)^3 |k_z|) ]', 'S14')

add_body(
    '对传播模（k_perp < k_0），k_z 为实数；'
    '对倏逝模（k_perp > k_0），|k_z| = kappa 为纯虚数的模。'
    '分母中的 |k_z| 意味着：单个高 k_perp 倏逝模式携带更少的真空涨落振幅。'
    '这一性质在物理上是自洽的——如果每个倏逝模式的振幅不衰减，倏逝场的总能量将发散。'
)

# IV.B
doc.add_heading('IV.B 相互作用哈密顿量：偶极子如何驱动量子场', level=2)

add_body(
    '到目前为止处理的全是自由场（无源）。但小孔衍射的物理本质是：'
    '位于小孔处的等效磁偶极子 M(t) 和电偶极子 P(t)'
    '作为源，驱动量子化电磁场从真空态演化为含有真实光子的态。'
    '这个驱动过程由相互作用哈密顿量描述（正文 Eq. 37）：'
)

add_eq(
    'H_hat_int(t) = -M(t) . (curl A_hat)|_{r=0} + P(t) . (dA_hat/dt)|_{r=0}',
    'S15'
)

add_body(
    '物理含义直截了当：第一项是磁偶极矩 M 与磁场算符 B_hat = curl A 的耦合能量，'
    '第二项是电偶极矩 P 与电场算符 E_hat = -dA/dt 的耦合能量。'
    '两者均取在原点 r=0 处，因为 Bethe 理论把小孔等效为一个点偶极子。'
)

add_body(
    '将量子化场算符代入 H_int，提取产生单光子的跃迁矩阵元 V_{k_perp, lam}（正文 Eq. 38），'
    '其中包含了 k x M_0（磁偶极贡献）和 omega_k P_0_perp（电偶极贡献）两个矢量结构因子。'
)

# IV.C
doc.add_heading('IV.C 海森堡运动方程与相干态振幅', level=2)

add_body(
    '有了总哈密顿量 H_hat = H_hat_0 + H_hat_int，量子模式的动力学演化由海森堡运动方程描述'
    '（正文 Eq. 39）：'
)

add_eq('i hbar d/dt a_hat_k(t) = [a_hat_k(t), H_hat]', 'S16')

add_body(
    '这里出现了一个微妙的技术要点。由于我们要处理的是单频源的稳态响应，'
    '直觉上似乎可以直接在二维 Weyl 角谱上操作。'
    '但海森堡方程的时间积分本质上依赖于连续的频率谱（极点留数积分），'
    '而二维角谱已经把频率固定在 omega_0。'
    '因此，求解动力学方程时必须先退回到未经降维的三维连续模式空间（正文 Eq. 6），'
    '完成积分后再映射回二维角谱。'
)

add_body(
    '在三维表象下，将总哈密顿量展开并利用玻色对易关系完成代数化简。'
    '自由场部分产生相位演化因子 e^{-i omega_k t}，'
    '相互作用部分提取出含固定频率驱动的源项。'
    '方程退化为一个一阶常微分方程（正文 Eq. 41）。'
    '引入绝热开关参数 eta -> 0+ 消除初始时刻的非物理发散后，'
    '积分给出稳态相干态振幅（正文 Eq. 43）：'
)

add_eq(
    'alpha_3D(k, lam) = -V_3D(k, lam) / [hbar (omega_k - omega_0 - i eta)]',
    'S17'
)

add_body(
    '这就是流程图中"Coherent State Amplitude <a_mu(t)>"节点的具体内容。'
    '它描述了每个模式 (k, lam) 在稳态下的平均光子数振幅。'
    '分母中的 i eta 给出微小的线宽，使得极点积分可以用主值和反常部分的标准方法处理。'
)

add_body(
    '最终，将三维相干态振幅映射回二维角谱（正文 Eq. 45），'
    '得到矢量势算符在相干态中的期望值。'
    '这个映射利用了雅可比行列式 dk_z = (omega_k / c^2 k_z) d omega_k 执行积分微元代换，'
    '将 k_z 积分转化为频率极点处的留数贡献。'
)

# ═══════════════════════════════════════════════════════════════
# V. RESULTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('V. 结果：经典极限与形状因子', level=1)
add_body('【对应流程图：右分支，从相干态振幅到 Bethe 结果的复现】')

# V.A
doc.add_heading('V.A 经典极限：量子结果退化为 Bethe 理论', level=2)

add_body(
    '任何量子理论都必须在适当极限下回归经典物理。'
    '对于电磁场量子化，这个极限对应于相干态的大振幅行为——'
    '当相干态的平均光子数足够大时，场算符的期望值应当精确等于经典场。'
)

add_body(
    '正文 Eq. 46 给出 Weyl 展开后矢量势算符在相干态中的期望值。'
    '将其与从三维动力学映射回的表达式（正文 Eq. 45）对比，'
    '可以解出二维相干态振幅 alpha_{k_perp, lam} 与跃迁矩阵元之间的关系（正文 Eq. 47）。'
    '代入 C_{k_perp} 的具体表达式，并对极化求和（利用完备性关系 S5），最终得到（正文 Eq. 48）：'
)

add_eq(
    '<A_hat(r, t)> = -(mu_0/8 pi^2) int d^2k_perp (k x M_0 - omega_0 P_0_perp)/k_z '
    '* e^{-i(omega_0 t - k_perp.rho - k_z z)}',
    'S18'
)

add_body(
    '这个二维积分正是 Sommerfeld 恒等式的逆向运用。'
    '当 k_perp 积分遍及整个平面时，结果等价于球面波的旋度运算，'
    '直接给出经典偶极辐射叠加（正文 Eq. 49）：'
)

add_eq(
    '<A_hat(r, t)> = (mu_0/4pi) curl[M_0 e^{-i(omega_0 t - k_0 r)}/r] '
    '- (i mu_0 omega_0/4pi) P_0_perp e^{-i(omega_0 t - k_0 r)}/r',
    'S19'
)

add_body(
    '第一项是磁偶极辐射的矢量势，第二项是电偶极辐射的矢量势。'
    '这正是经典电动力学教材中的标准结果。'
    '至此，流程图中"Classical limit"节点的内容得到了严格验证。'
)

# V.B
doc.add_heading('V.B 圆形小孔的形状因子：从点源到有限尺寸', level=2)

add_body(
    '以上所有推导都假设偶极子位于原点——一个数学上的点。但真实的小孔有有限的半径 a。'
    'Bethe 早在 1944 年就指出，小孔面内的感应电流不是集中在中心点的，'
    '而是按照一个椭圆分布展宽（正文 Eq. 52）：'
)

add_eq('m(rho) proportional to sqrt(1 - (rho/a)^2),   rho <= a', 'S20')

add_body(
    '这个分布在边缘 rho = a 处趋零，在中心 rho = 0 处取最大值，'
    '物理上反映了孔边缘电流被截断的边界条件。'
    '为了在角谱框架中纳入有限尺寸效应，'
    '需要计算这个面电流分布的二维 Fourier 变换（正文 Eq. 53-54）。'
)

add_body(
    '变量替换 u = rho/a 后，积分化为 Sonine 有限积分（正文 Eq. 55-56）。'
    '利用标准结果 int_0^1 u sqrt(1-u^2) J_0(cu) du = j_1(c)/c，'
    '其中 j_1 是第一类球 Bessel 函数，得到（正文 Eq. 57）：'
)

add_eq('m_tilde(k_perp) = 3(M_0 + P_0) * j_1(k_perp a) / (k_perp a)', 'S21')

add_body(
    '形状因子 j_1(k_perp a)/(k_perp a) 的物理含义非常直观：'
    '当 k_perp a << 1（低空间频率，即远场），j_1(x)/x 趋近 1/3，'
    '形状因子趋近常数——孔的有限尺寸可忽略，点偶极近似成立。'
    '当 k_perp a >> 1（高空间频率，即极端近场），j_1(x)/x 近似 1/(k_perp a)^2 快速衰减——'
    '高波矢分量被小孔的有限尺寸自然截断。'
)

add_body(
    '这个衰减机制在理论层面解释了一个重要问题：'
    '即使倏逝模式的归一化振幅允许 k_perp -> 无穷大 的模式存在，'
    '形状因子 j_1(k_perp a)/(k_perp a) 为极端近场提供了一个内禀的紫外截断。'
    '因此场的量子期望值在任意近的距离上始终有限且平滑收敛，不存在物理奇异性。'
)

# V.C
doc.add_heading('V.C 完整结果与 Bethe 截面', level=2)

add_body(
    '将形状因子纳入相干态振幅，得到修正后的跃迁矩阵元（正文 Eq. 58）和完整的矢量势期望值（正文 Eq. 59）：'
)

add_eq(
    '<A_hat(r, t)> = -(mu_0/8 pi^2) int d^2k_perp (k x M_0 - omega_0 P_0)/k_z '
    '* 3 j_1(k_perp a)/(k_perp a) * e^{-i(omega_0 t - k_perp.rho - k_z z)}',
    'S22'
)

add_body(
    '这是全文的核心结果。它给出了衍射场在距小孔任意距离处'
    '（包括倏逝模式主导的近场区 z 约小于 a）的量子平均矢量势。'
    '在远场（k_0 r >> 1, a << r），当 k_perp a << 1 时 j_1(k_perp a)/(k_perp a) 趋近 1/3，'
    '上式退化为 (S19)，即点孔情形。'
)

add_body(
    '由矢量势推导辐射功率，对远场半球积分，'
    '代入 Bethe 的等效偶极矩 M_0 = (8/3) a^3 H_0 和 P_0 = -(4/3) a^3 E_0，'
    '最终给出总辐射截面（正文 Eq. 62）：'
)

add_eq('sigma_t = (64/27 pi) k_0^4 a^6', 'S23')

add_body(
    '这正是 Bethe 1944 年得到的经典结果。'
    '流程图中"Reproduction of Bethe Results"节点的验证至此全部完成。'
    '量子化理论不仅在经典极限下严格回归 Bethe 理论，'
    '还通过形状因子拓展到了有限孔径的近场区域——'
    '这是经典理论无法直接处理的前沿。'
)

# ═══════════════════════════════════════════════════════════════
# VI. SUMMARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('VI. 推导链条总结', level=1)

add_body(
    '回顾整条推导链路，可以将其浓缩为以下逻辑骨架：'
)

add_body(
    '(1) 起点：经典小孔衍射的 Bethe 偶极模型（M_0, P_0）给出了物理直觉——'
    '小孔等效为点偶极子。但经典理论无法描述光子统计和真空涨落。'
)

add_body(
    '(2) 量子化基础：库仑规范下，从拉格朗日量 L 出发，定义正则共轭动量 Pi = -eps_0 E，'
    '通过等时对易关系和模展开确定归一化系数 N_k，得到完整的场算符 E_hat、B_hat 和自由哈密顿量 H_hat_0。'
)

add_body(
    '(3) 协变验证：洛伦兹规范下，四分量 Fermi 拉格朗日量引入了非物理自由度，'
    'Gupta-Bleuler 条件剔除了这些自由度，最终证明两种规范给出相同的物理可观测量——'
    '结果与规范选取无关。'
)

add_body(
    '(4) 动力学求解：Weyl 角谱展开将三维积分降为二维，'
    '相互作用哈密顿量描述偶极子对量子场的驱动，'
    '海森堡运动方程给出稳态相干态振幅。'
)

add_body(
    '(5) 经典验证：相干态期望值在经典极限下精确复现 Bethe 辐射矢量势和截面。'
    '有限孔径引入球 Bessel 形状因子，为近场提供物理截断。'
)

add_body(
    '五个环节环环相扣，每一步都有明确的物理动机和数学必要性。'
    '去掉任何一环，链条都会断裂——这就是这套量子化框架的自洽性所在。'
)

# ── Save ──
output_path = r'z:\321\DHL\Self_Learning\DHL\small_hole_qed\Hole_quantization\Supplemental_Material.docx'
doc.save(output_path)
print('Saved to: %s' % output_path)
print('Done.')
